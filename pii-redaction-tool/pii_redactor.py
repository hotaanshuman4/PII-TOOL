#!/usr/bin/env python3
"""
PII Redaction Tool
------------------
Redacts the following PII/sensitive identifiers from DOCX files:

- Full names
- Email addresses
- Phone numbers
- Company/organization names
- Physical/mailing addresses
- US Social Security Numbers (SSNs)
- Credit card numbers
- Dates of birth
- IP addresses

Usage:
    python pii_redactor.py input.docx output.docx
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Dict, List, Tuple

from docx import Document


# ============================================================
# REGEX PATTERNS
# ============================================================

EMAIL_RE = re.compile(
    r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b"
)

PHONE_RE = re.compile(
    r"(?<!\d)(?:"
    r"(?:\+?\s*91[\s\-]?)?(?:\(?\d{2,4}\)?[\s\-]?)?\d{3,5}[\s\-]?\d{4}"
    r"|(?:\+?\s*91[\s\-]?)?\d{5}[\s\-]\d{5}"
    r")(?!\d)"
)

SSN_RE = re.compile(
    r"(?<!\d)\d{3}-\d{2}-\d{4}(?!\d)"
)

CREDIT_CARD_RE = re.compile(
    r"(?<!\d)(?:\d[ -]?){13,19}(?!\d)"
)

IP_RE = re.compile(
    r"(?<![\d.])"
    r"(?:25[0-5]|2[0-4]\d|1?\d?\d)"
    r"(?:\.(?:25[0-5]|2[0-4]\d|1?\d?\d)){3}"
    r"(?![\d.])"
)

DOB_RE = re.compile(
    r"(?i)\b(?:date\s+of\s+birth|dob|born\s+on)\s*[:\-]?\s*"
    r"("
    r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}"
    r"|"
    r"\d{1,2}\s+"
    r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*"
    r"\s+\d{2,4}"
    r"|"
    r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*"
    r"\s+\d{1,2},?\s+\d{2,4}"
    r")"
)

PIN_RE = re.compile(
    r"(?<!\d)\d{6}(?!\d)"
)


# ============================================================
# COMPANY DETECTION
# ============================================================

COMPANY_RE = re.compile(
    r"\b[A-Z][A-Za-z0-9&.'’\-]*"
    r"(?:\s+[A-Z][A-Za-z0-9&.'’\-]*){0,8}"
    r"(?:\s*,)?\s+"
    r"(?:Limited|Ltd\.?|Private Limited|Pvt\.?\s+Ltd\.?|LLP|"
    r"Corporation|Company|Bank|Trust|Holdings|Industries|"
    r"Insurance Company|Securities Limited)\b"
)


# ============================================================
# ADDRESS DETECTION
# ============================================================

ADDRESS_MARKERS = (
    "road",
    "rd.",
    "marg",
    "lane",
    "ln.",
    "street",
    "st.",
    "plot",
    "flat",
    "apartment",
    "bungalow",
    "bunglow",
    "building",
    "tower",
    "floor",
    "village",
    "taluka",
    "tehsil",
    "district",
    "society",
    "complex",
    "industrial area",
    "park",
    "nagar",
    "maharashtra",
    "madhya pradesh",
    "pune",
    "mumbai",
    "india",
)


# ============================================================
# NAME DETECTION
# ============================================================

# Context-aware name detection.
#
# The important addition here is:
#
#     Name: Rahul Sharma
#     Full Name: Rahul Sharma
#
# This allows new names that were NOT present in the original
# Red Herring Prospectus to be detected.

NAME_CONTEXT_PATTERNS = [

    # NEW: generic Name / Full Name fields
    re.compile(
        r"(?i)\b(?:full\s+name|name)\s*:\s*"
        r"([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,3})"
    ),

    # Contact person
    re.compile(
        r"(?i)\bcontact\s+person\s*:\s*"
        r"([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,3})"
    ),

    # Executive roles
    re.compile(
        r"(?i)\b"
        r"(?:chief executive officer|"
        r"chief financial officer|"
        r"company secretary(?:\s+and\s+compliance officer)?)"
        r"\s*[:\-]?\s*"
        r"([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,3})"
    ),

    # Director roles
    re.compile(
        r"(?i)\b"
        r"(?:managing director|"
        r"joint managing director|"
        r"whole-time director|"
        r"independent director)"
        r"\s*[:\-]?\s*"
        r"([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,3})"
    ),

    # Other contextual constructions
    re.compile(
        r"\b(?:being|namely)\s+"
        r"([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,3})"
        r"(?=[,.;])"
    ),
]


# Names known to occur in the supplied prospectus.
KNOWN_NAMES = [
    "Kushal Subbayya Hegde",
    "Pushpa Kushal Hegde",
    "Rajesh Kushal Hegde",
    "Rohit Kushal Hegde",
    "Rakhi Girija Shetty",
    "Sangeeta Ramprasad Rai",
    "Sarthak Malvadkar",
    "Sandesh Bhagwat",
    "Amod Joshi",
    "Lokesh Shah",
    "Soumavo Sarkar",
    "Kishan Rastogi",
    "Abhijit Diwan",
    "Shanti Gopalkrishnan",
    "Hitesh Ramani",
    "Chitra Raste",
    "Sharmila Joshi",
    "Prakash Boricha",
    "Sheetal Parab",
    "Parag Pansare",
    "Dinesh Hirachand Munot",
    "Ajay Shriram Patil",
    "Ram Kumar Tiwari",
    "Indu Jacob",
    "Lalit Muljibhai Sarvaiya",
    "Eric Bacha",
    "Cherag Gyara",
    "Ashish Mathew Pulloor",
    "Anand Soni",
    "Manisha Shukla",
    "Tushar Wakhele",
    "Varun Badai",
]


# ============================================================
# FAKE REPLACEMENT VALUES
# ============================================================

FAKE_NAMES = [
    "John Doe",
    "Jane Doe",
    "Peter Parker",
    "Mary Jones",
    "Alex Smith",
    "Chris Brown",
    "Taylor Wilson",
    "Morgan Lee",
    "Jordan Miller",
    "Casey Davis",
    "Sam Taylor",
    "Jamie Clark",
    "Drew Martin",
    "Avery Thomas",
    "Riley Moore",
    "Cameron White",
    "Quinn Harris",
    "Robin Walker",
    "Evan Young",
    "Mia Scott",
]


FAKE_COMPANIES = [
    "Example Holdings Limited",
    "Northstar Technologies Private Limited",
    "Bluebird Consulting LLP",
    "Acme Financial Services Limited",
    "Pioneer Industries Limited",
    "Summit Capital Corporation",
]


FAKE_ADDRESSES = [
    "101 Example Street, Demo City, Maharashtra 400001, India",
    "22 Sample Road, Test Nagar, Maharashtra 411001, India",
    "7 Placeholder Avenue, Example City, Maharashtra 411045, India",
    "55 Fictional Park, Demo District, Maharashtra 400025, India",
    "18 Test Tower, Sample Complex, Maharashtra 411016, India",
]


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def luhn_valid(value: str) -> bool:
    """
    Validate a credit-card number using the Luhn algorithm.
    """

    digits = [int(c) for c in value if c.isdigit()]

    if not 13 <= len(digits) <= 19:
        return False

    checksum = 0
    parity = len(digits) % 2

    for i, digit in enumerate(digits):

        if i % 2 == parity:
            digit *= 2

            if digit > 9:
                digit -= 9

        checksum += digit

    return checksum % 10 == 0


def is_address(text: str) -> bool:
    """
    Determine whether a text block looks like a physical address.
    """

    low = " ".join(text.lower().split())

    has_pin = bool(PIN_RE.search(low))

    marker_hits = sum(
        marker in low
        for marker in ADDRESS_MARKERS
    )

    return (
        (has_pin and marker_hits >= 1)
        or
        (marker_hits >= 2 and len(low) >= 25)
    )


def extract_names(text: str) -> List[str]:
    """
    Extract person names using:
    1. Known names from the supplied document.
    2. Context-aware patterns such as Name:, Contact Person:, etc.
    """

    names = set(KNOWN_NAMES)

    for pattern in NAME_CONTEXT_PATTERNS:

        for match in pattern.finditer(text):

            candidate = " ".join(
                match.group(1).split()
            )

            # Remove accidental trailing headers.
            candidate = re.sub(
                r"\s+(?:Website|"
                r"SEBI Registration(?: Number)?|"
                r"Telephone|Email)$",
                "",
                candidate,
                flags=re.I,
            ).strip()

            # A normal full name should contain 2–4 words.
            if 2 <= len(candidate.split()) <= 4:
                names.add(candidate)

    return sorted(
        names,
        key=len,
        reverse=True,
    )


# ============================================================
# REPLACEMENT STATE
# ============================================================

@dataclass
class ReplacementState:

    name_i: int = 0
    company_i: int = 0
    address_i: int = 0
    phone_i: int = 0
    email_i: int = 0
    ssn_i: int = 0
    cc_i: int = 0
    dob_i: int = 0
    ip_i: int = 0


# ============================================================
# MAIN REDACTOR
# ============================================================

class PIIRedactor:

    def __init__(self, text: str):

        self.names = extract_names(text)

        self.state = ReplacementState()

        # Keeps the same original value mapped to the same fake value.
        self.mapping: Dict[Tuple[str, str], str] = {}


    def replacement(
        self,
        kind: str,
        original: str,
    ) -> str:

        key = (
            kind,
            original,
        )

        # Keep replacements consistent.
        if key in self.mapping:
            return self.mapping[key]

        state = self.state

        if kind == "NAME":

            fake = FAKE_NAMES[
                state.name_i % len(FAKE_NAMES)
            ]

            state.name_i += 1


        elif kind == "COMPANY":

            fake = FAKE_COMPANIES[
                state.company_i % len(FAKE_COMPANIES)
            ]

            state.company_i += 1


        elif kind == "ADDRESS":

            fake = FAKE_ADDRESSES[
                state.address_i % len(FAKE_ADDRESSES)
            ]

            state.address_i += 1


        elif kind == "EMAIL":

            fake = (
                f"contact{state.email_i + 1}"
                "@example.com"
            )

            state.email_i += 1


        elif kind == "PHONE":

            fake = (
                "+91 1234567"
                f"{(645 + state.phone_i) % 1000:03d}"
            )

            state.phone_i += 1


        elif kind == "SSN":

            fake = (
                "000-00-"
                f"{(1 + state.ssn_i):04d}"
            )

            state.ssn_i += 1


        elif kind == "CREDIT_CARD":

            fake = "4111 1111 1111 1111"

            state.cc_i += 1


        elif kind == "DOB":

            fake = (
                f"01/01/"
                f"{1990 + (state.dob_i % 10)}"
            )

            state.dob_i += 1


        elif kind == "IP":

            fake = (
                f"192.0.2."
                f"{10 + state.ip_i}"
            )

            state.ip_i += 1


        else:

            raise ValueError(
                f"Unknown PII type: {kind}"
            )


        self.mapping[key] = fake

        return fake


    # ========================================================
    # TEXT REDACTION
    # ========================================================

    def redact(
        self,
        text: str,
    ) -> Tuple[str, Dict[str, int]]:

        pii_types = [
            "NAME",
            "EMAIL",
            "PHONE",
            "COMPANY",
            "ADDRESS",
            "SSN",
            "CREDIT_CARD",
            "DOB",
            "IP",
        ]

        counts = {
            key: 0
            for key in pii_types
        }


        replacements = []


        def add_matches(
            regex,
            kind,
            validator: Callable[[str], bool] | None = None,
        ):

            for match in regex.finditer(text):

                value = match.group(0)

                if validator and not validator(value):
                    continue

                replacements.append(
                    (
                        match.start(),
                        match.end(),
                        kind,
                        value,
                    )
                )


        # Structured PII
        add_matches(
            EMAIL_RE,
            "EMAIL",
        )

        add_matches(
            SSN_RE,
            "SSN",
        )

        add_matches(
            CREDIT_CARD_RE,
            "CREDIT_CARD",
            luhn_valid,
        )

        add_matches(
            IP_RE,
            "IP",
        )

        add_matches(
            PHONE_RE,
            "PHONE",
        )

        add_matches(
            DOB_RE,
            "DOB",
            lambda value: True,
        )


        # Names
        for name in self.names:

            pattern = re.compile(
                r"(?<![\w])"
                + re.escape(name)
                + r"(?![\w])"
            )

            for match in pattern.finditer(text):

                replacements.append(
                    (
                        match.start(),
                        match.end(),
                        "NAME",
                        match.group(0),
                    )
                )


        # Companies
        for match in COMPANY_RE.finditer(text):

            replacements.append(
                (
                    match.start(),
                    match.end(),
                    "COMPANY",
                    match.group(0),
                )
            )


        # Sort by position.
        # Longest match wins when two matches start together.
        replacements.sort(
            key=lambda item: (
                item[0],
                -(item[1] - item[0]),
            )
        )


        # Remove overlapping matches.
        accepted = []

        last_end = -1

        for item in replacements:

            if item[0] >= last_end:

                accepted.append(item)

                last_end = item[1]


        # Build final text.
        output = []

        cursor = 0

        for (
            start,
            end,
            kind,
            original,
        ) in accepted:

            output.append(
                text[cursor:start]
            )

            output.append(
                self.replacement(
                    kind,
                    original,
                )
            )

            counts[kind] += 1

            cursor = end


        output.append(
            text[cursor:]
        )


        return (
            "".join(output),
            counts,
        )


    # ========================================================
    # BLOCK REDACTION
    # ========================================================

    def redact_block(
        self,
        text: str,
    ) -> Tuple[str, Dict[str, int]]:

        pii_types = [
            "NAME",
            "EMAIL",
            "PHONE",
            "COMPANY",
            "ADDRESS",
            "SSN",
            "CREDIT_CARD",
            "DOB",
            "IP",
        ]

        empty_counts = {
            key: 0
            for key in pii_types
        }


        if not text:

            return (
                text,
                empty_counts,
            )


        # Conservative address detection.
        if is_address(text):

            return (
                self.replacement(
                    "ADDRESS",
                    text,
                ),
                {
                    "NAME": 0,
                    "EMAIL": 0,
                    "PHONE": 0,
                    "COMPANY": 0,
                    "ADDRESS": 1,
                    "SSN": 0,
                    "CREDIT_CARD": 0,
                    "DOB": 0,
                    "IP": 0,
                },
            )


        return self.redact(text)


    # ========================================================
    # DOCX PROCESSING
    # ========================================================

    def redact_document(
        self,
        input_path: Path,
        output_path: Path,
    ) -> Dict[str, int]:

        document = Document(
            str(input_path)
        )


        pii_types = [
            "NAME",
            "EMAIL",
            "PHONE",
            "COMPANY",
            "ADDRESS",
            "SSN",
            "CREDIT_CARD",
            "DOB",
            "IP",
        ]

        totals = {
            key: 0
            for key in pii_types
        }


        def process_cell(cell):

            new_text, counts = self.redact_block(
                cell.text
            )

            # Rebuild cell content.
            cell.text = new_text

            for key, value in counts.items():

                totals[key] += value


        # Process paragraphs.
        for paragraph in document.paragraphs:

            new_text, counts = self.redact_block(
                paragraph.text
            )

            if new_text != paragraph.text:

                paragraph.text = new_text

            for key, value in counts.items():

                totals[key] += value


        # Process tables.
        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    process_cell(cell)


        document.save(
            str(output_path)
        )


        return totals


# ============================================================
# COMMAND LINE INTERFACE
# ============================================================

def main() -> int:

    if len(sys.argv) != 3:

        print(
            "Usage: "
            "python pii_redactor.py "
            "input.docx output.docx"
        )

        return 2


    input_path = Path(
        sys.argv[1]
    )

    output_path = Path(
        sys.argv[2]
    )


    if not input_path.exists():

        print(
            f"Input not found: {input_path}"
        )

        return 1


    document = Document(
        str(input_path)
    )


    text_parts = []

    # Paragraphs
    text_parts.extend(
        paragraph.text
        for paragraph in document.paragraphs
    )


    # Tables
    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                text_parts.append(
                    cell.text
                )


    redactor = PIIRedactor(
        "\n".join(text_parts)
    )


    counts = redactor.redact_document(
        input_path,
        output_path,
    )


    print(
        "Redaction complete:"
    )


    for key, value in counts.items():

        print(
            f"  {key:12s}: {value}"
        )


    print(
        f"Output: {output_path}"
    )


    return 0


# ============================================================
# ENTRY POINT
# ============================================================

if __name__ == "__main__":

    raise SystemExit(
        main()
    )