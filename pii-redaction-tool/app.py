import io
import tempfile
from pathlib import Path

import streamlit as st
from docx import Document

from pii_redactor import PIIRedactor

st.set_page_config(
    page_title="PII Redaction Tool",
    page_icon="🔐",
    layout="centered",
)

st.title("🔐 PII Redaction Tool")
st.write(
    "Upload a DOCX document to detect and replace supported personally "
    "identifiable information (PII) with deterministic fake alternatives."
)

st.info(
    "Supported: full names, email addresses, phone numbers, company/"
    "organization names, physical/mailing addresses, SSNs, credit cards, "
    "dates of birth, and IP addresses."
)

uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])

if uploaded_file is not None:
    if st.button("Redact PII", type="primary"):
        temp_input = None
        temp_output = None

        try:
            # Save the upload to temporary files because the core engine's
            # public API accepts input/output paths.
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
                f.write(uploaded_file.getvalue())
                temp_input = Path(f.name)

            temp_output = temp_input.with_name(temp_input.stem + "_redacted.docx")

            # Read all document text first so context-aware name detection
            # sees paragraphs and table cells from the whole document.
            doc = Document(str(temp_input))
            text_parts = [p.text for p in doc.paragraphs]

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

            redactor = PIIRedactor("\n".join(text_parts))
            counts = redactor.redact_document(temp_input, temp_output)

            with open(temp_output, "rb") as f:
                output_bytes = f.read()

            st.success("Redaction completed successfully.")

            st.subheader("Detection summary")

            labels = {
                "NAME": "Names",
                "EMAIL": "Emails",
                "PHONE": "Phone Numbers",
                "COMPANY": "Companies",
                "ADDRESS": "Addresses",
                "SSN": "SSNs",
                "CREDIT_CARD": "Credit Cards",
                "DOB": "Dates of Birth",
                "IP": "IP Addresses",
            }

            total = sum(counts.values())
            cols = st.columns(3)

            for i, key in enumerate(labels):
                cols[i % 3].metric(labels[key], counts.get(key, 0))

            st.metric("Total replacements", total)

            output_name = f"{Path(uploaded_file.name).stem}_redacted.docx"

            st.download_button(
                "⬇️ Download Redacted DOCX",
                data=output_bytes,
                file_name=output_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
            )

        except Exception as exc:
            st.error(f"Could not process the document: {exc}")
            st.exception(exc)

        finally:
            for path in (temp_input, temp_output):
                if path is not None:
                    try:
                        path.unlink(missing_ok=True)
                    except Exception:
                        pass

st.divider()
st.caption(
    "Documents are processed in the current application session. "
    "The app does not intentionally persist uploaded files."
)
